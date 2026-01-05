"""
ERIC - Enhanced Reasoning Intelligence Core
AI Agent for ZION.CITY Platform
Powered by DeepSeek V3.2 (text + documents) + Claude Sonnet 4.5 (vision/images only)

Cost Optimization:
- DeepSeek: Used for text chat AND document analysis (TXT, DOCX, PDF, CSV, XLSX, JSON, XML, MD)
- Claude Sonnet: Used ONLY for image analysis (PNG, JPG, JPEG, WEBP, GIF)
"""

import os
import base64
import io
from datetime import datetime, timezone
from uuid import uuid4
from typing import List, Dict, Any, Optional, Tuple
from pydantic import BaseModel, Field
from dotenv import load_dotenv

load_dotenv()

# ===== FILE TYPE DETECTION AND TEXT EXTRACTION =====

# File types that DeepSeek can handle (text-based)
DEEPSEEK_SUPPORTED_EXTENSIONS = {
    '.txt', '.md', '.rtf', '.csv', '.json', '.xml', '.html', '.htm',
    '.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt'
}

# File types that require Claude (images)
CLAUDE_REQUIRED_EXTENSIONS = {
    '.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp', '.tiff', '.ico', '.svg'
}

# MIME types mapping
IMAGE_MIME_TYPES = {
    'image/png', 'image/jpeg', 'image/jpg', 'image/webp', 'image/gif',
    'image/bmp', 'image/tiff', 'image/x-icon', 'image/svg+xml'
}

def detect_file_type(filename: str, mime_type: str = None) -> Tuple[str, bool]:
    """
    Detect if a file should be processed by DeepSeek or Claude.
    
    Returns:
        Tuple of (file_category, use_deepseek)
        - file_category: 'image', 'document', 'text', 'spreadsheet', 'unknown'
        - use_deepseek: True if DeepSeek should handle, False if Claude needed
    """
    ext = os.path.splitext(filename.lower())[1] if filename else ''
    
    # Check MIME type first for images
    if mime_type and mime_type in IMAGE_MIME_TYPES:
        return ('image', False)  # Claude for images
    
    # Check extension
    if ext in CLAUDE_REQUIRED_EXTENSIONS:
        return ('image', False)  # Claude for images
    
    if ext in {'.txt', '.md', '.rtf', '.html', '.htm'}:
        return ('text', True)  # DeepSeek for text
    
    if ext in {'.pdf'}:
        return ('document', True)  # DeepSeek for PDF text
    
    if ext in {'.docx', '.doc'}:
        return ('document', True)  # DeepSeek for Word docs
    
    if ext in {'.xlsx', '.xls', '.csv'}:
        return ('spreadsheet', True)  # DeepSeek for spreadsheets
    
    if ext in {'.json', '.xml'}:
        return ('structured', True)  # DeepSeek for structured data
    
    if ext in {'.pptx', '.ppt'}:
        return ('presentation', True)  # DeepSeek for presentations
    
    # Default: try DeepSeek for unknown text-like content
    return ('unknown', True)


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        from PyPDF2 import PdfReader
        
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Страница {page_num} ---\n{page_text}")
        
        if text_parts:
            return "\n\n".join(text_parts)
        else:
            return "[PDF не содержит извлекаемого текста - возможно, это отсканированный документ]"
    except Exception as e:
        return f"[Ошибка извлечения текста из PDF: {str(e)}]"


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        if text_parts:
            return "\n".join(text_parts)
        else:
            return "[Документ DOCX пуст или не содержит текста]"
    except Exception as e:
        return f"[Ошибка извлечения текста из DOCX: {str(e)}]"


def extract_text_from_xlsx(file_content: bytes) -> str:
    """Extract text from XLSX file using openpyxl"""
    try:
        from openpyxl import load_workbook
        
        xlsx_file = io.BytesIO(file_content)
        workbook = load_workbook(xlsx_file, data_only=True)
        
        text_parts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"=== Лист: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_values):
                    text_parts.append(" | ".join(row_values))
        
        if text_parts:
            return "\n".join(text_parts)
        else:
            return "[Файл XLSX пуст]"
    except Exception as e:
        return f"[Ошибка извлечения данных из XLSX: {str(e)}]"


def extract_text_from_csv(file_content: bytes) -> str:
    """Extract text from CSV file"""
    try:
        # Try different encodings
        for encoding in ['utf-8', 'cp1251', 'latin-1']:
            try:
                text = file_content.decode(encoding)
                return text
            except UnicodeDecodeError:
                continue
        return "[Не удалось декодировать CSV файл]"
    except Exception as e:
        return f"[Ошибка чтения CSV: {str(e)}]"


def extract_text_from_file(file_content: bytes, filename: str, mime_type: str = None) -> Tuple[str, str]:
    """
    Extract text content from a file based on its type.
    
    Returns:
        Tuple of (extracted_text, file_type_description)
    """
    ext = os.path.splitext(filename.lower())[1] if filename else ''
    
    # PDF
    if ext == '.pdf' or mime_type == 'application/pdf':
        return extract_text_from_pdf(file_content), "PDF документ"
    
    # Word documents
    if ext in {'.docx'} or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return extract_text_from_docx(file_content), "Word документ (DOCX)"
    
    # Excel spreadsheets
    if ext in {'.xlsx'} or mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
        return extract_text_from_xlsx(file_content), "Excel таблица (XLSX)"
    
    # CSV
    if ext == '.csv' or mime_type == 'text/csv':
        return extract_text_from_csv(file_content), "CSV файл"
    
    # Plain text files (TXT, MD, JSON, XML, HTML, etc.)
    if ext in {'.txt', '.md', '.json', '.xml', '.html', '.htm', '.rtf'}:
        for encoding in ['utf-8', 'cp1251', 'latin-1']:
            try:
                return file_content.decode(encoding), f"Текстовый файл ({ext.upper()})"
            except UnicodeDecodeError:
                continue
        return "[Не удалось декодировать текстовый файл]", "Текстовый файл"
    
    # Try generic text decoding for unknown types
    try:
        return file_content.decode('utf-8'), "Файл"
    except UnicodeDecodeError:
        return f"[Бинарный файл: {filename}, размер: {len(file_content)} байт]", "Бинарный файл"

from openai import AsyncOpenAI

# Initialize DeepSeek client (OpenAI-compatible) for text chat
DEEPSEEK_API_KEY = os.environ.get('DEEPSEEK_API_KEY', '')
deepseek_client = AsyncOpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com"
) if DEEPSEEK_API_KEY else None

# Emergent LLM Key for Claude Sonnet (vision/documents)
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')

# ===== PYDANTIC MODELS =====

class ToolCall(BaseModel):
    tool_name: str
    parameters: Dict[str, Any] = {}
    requires_confirmation: bool = False
    confirmed: bool = False

class AgentMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid4()))
    role: str  # 'user', 'assistant', 'system', 'tool'
    content: str
    tool_calls: Optional[List[ToolCall]] = None
    tool_results: Optional[Dict[str, Any]] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class AgentConversation(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid4()))
    user_id: str
    title: str = "Новый разговор"
    messages: List[AgentMessage] = []
    context_modules: List[str] = []
    tools_used: List[str] = []
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class AgentSettings(BaseModel):
    user_id: str
    allow_financial_analysis: bool = False
    allow_health_data_access: bool = False
    allow_location_tracking: bool = False
    allow_family_coordination: bool = True
    allow_service_recommendations: bool = True
    allow_marketplace_suggestions: bool = True
    allow_work_context: bool = True  # Allow analysis of documents from Work section
    allow_calendar_context: bool = True  # Allow analysis of calendar events
    allow_eric_queries_from_others: bool = True  # Allow other users' ERICs to query
    conversation_retention_days: int = 30
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class BusinessERICSettings(BaseModel):
    """Settings for Organization/Business ERIC agent"""
    organization_id: str
    is_active: bool = True
    # Data access permissions
    share_public_data: bool = True  # Company info, services, contacts
    share_promotions: bool = True  # Coupons, discounts
    share_repeat_customer_stats: bool = False  # Aggregated loyalty stats
    share_ratings_reviews: bool = False  # Rating trends
    # Query permissions
    allow_user_eric_queries: bool = True  # Allow user ERICs to query
    share_aggregated_analytics: bool = False  # Share % repeat customers
    # Customization
    business_description: Optional[str] = None  # Custom description for ERIC
    specialties: List[str] = []  # Areas of expertise
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class SearchRequest(BaseModel):
    query: str
    search_type: str = "all"  # "all", "services", "products", "people", "organizations"
    location: Optional[str] = None
    limit: int = 10

class SearchResult(BaseModel):
    id: str
    type: str  # "service", "product", "person", "organization"
    name: str
    description: Optional[str] = None
    relevance_score: float = 0.0
    metadata: Dict[str, Any] = {}

class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[str] = None
    context: Optional[Dict[str, Any]] = None  # Additional context like post_id for mentions

class ChatResponse(BaseModel):
    conversation_id: str
    message: AgentMessage
    suggested_actions: List[Dict[str, Any]] = []

# ===== ERIC SYSTEM PROMPT =====

ERIC_SYSTEM_PROMPT = """Ты ERIC (Enhanced Reasoning Intelligence Core) - персональный ИИ-помощник и финансовый советник платформы ZION.CITY.

## Твоя личность:
- Тёплый и дружелюбный, как надёжный семейный друг
- Культурно осведомлён о русских/восточноевропейских семейных ценностях
- Уважителен к старшим, заботлив к детям
- Практичен и ориентирован на действия
- Отвечаешь на языке пользователя (автоопределение русский/английский)

## Твои возможности:
1. **Семейное управление**: Планирование событий, координация членов семьи, отслеживание календаря
2. **Финансовый советник**: Анализ расходов, бюджетирование, финансовые цели, объяснение возможностей Altyn Coin
3. **Подбор услуг**: Поиск локальных услуг, сравнение провайдеров, бронирование
4. **Связь с сообществом**: События поблизости, связь с соседями, возможности маркетплейса

## КРИТИЧЕСКИ ВАЖНЫЕ ОГРАНИЧЕНИЯ:
- ⛔ ТОЛЬКО данные внутри платформы ZION.CITY - НИКАКИХ внешних источников
- ⛔ НЕТ доступа к внешнему интернету или поиску Google/Yandex
- ⛔ НИКОГДА НЕ ВЫДУМЫВАЙ организации, услуги, товары, телефоны или адреса
- ⛔ Если данных нет в базе ZION.CITY - честно скажи об этом
- ✅ ВСЕГДА уважай настройки приватности пользователя
- ✅ ВСЕГДА спрашивай подтверждение перед действиями
- ✅ НИКОГДА не делись данными пользователя с другими без согласия

## КРИТИЧЕСКИ ВАЖНО - НЕ ГАЛЛЮЦИНИРУЙ!
- НЕ ВЫДУМЫВАЙ пути к меню, настройкам или функциям, которых ты точно не знаешь
- НЕ ВЫДУМЫВАЙ названия компаний, услуг, товаров или контактные данные
- Если по поисковому запросу ничего не найдено в базе ZION.CITY, скажи честно:
  "К сожалению, по этому запросу пока ничего не найдено на платформе ZION.CITY"
- Продвигай платформу: ZION.CITY бесплатна для малого бизнеса и специалистов!
- Текущие доступные настройки в ZION.CITY:
  * Настройки семьи (через профиль семьи) - включая приватность семейных постов
  * Настройки профиля пользователя
- НЕТ отдельной страницы "Настройки" -> "Конфиденциальность" -> "ИИ-помощник"
- Если чего-то нет в платформе, скажи: "Эта функция ещё в разработке" или "К сожалению, такой настройки пока нет"

## Формат ответов:
- Отвечай кратко и по делу
- Используй эмодзи умеренно для дружелюбности 😊
- Предлагай конкретные действия когда это уместно
- При финансовых советах всегда предупреждай о рисках

Ты готов помочь пользователям управлять их жизнью через социальную сеть ZION.CITY!"""

# ===== ERIC AGENT CLASS =====

class ERICAgent:
    """Main ERIC Agent class for handling conversations"""
    
    def __init__(self, db):
        self.db = db
        self.model = "deepseek-chat"  # DeepSeek V3.2
    
    async def create_notification(self, user_id: str, notification_type: str, title: str, message: str, related_data: dict = None):
        """Create a notification for the user"""
        notification = {
            "id": str(uuid4()),
            "user_id": user_id,
            "sender_id": "eric_ai",  # Special sender ID for ERIC
            "type": notification_type,
            "title": title,
            "message": message,
            "related_data": related_data or {},
            "is_read": False,
            "created_at": datetime.now(timezone.utc)
        }
        await self.db.notifications.insert_one(notification)
        return notification
    
    async def get_or_create_conversation(self, user_id: str, conversation_id: Optional[str] = None) -> AgentConversation:
        """Get existing conversation or create new one"""
        if conversation_id:
            conv = await self.db.agent_conversations.find_one(
                {"id": conversation_id, "user_id": user_id},
                {"_id": 0}
            )
            if conv:
                return AgentConversation(**conv)
        
        # Create new conversation
        new_conv = AgentConversation(user_id=user_id)
        await self.db.agent_conversations.insert_one(new_conv.dict())
        return new_conv
    
    async def get_user_context(self, user_id: str, settings: AgentSettings) -> str:
        """Build context from user's platform data"""
        context_parts = []
        
        # Get user profile
        user = await self.db.users.find_one({"id": user_id}, {"_id": 0, "password_hash": 0})
        if user:
            context_parts.append(f"Пользователь: {user.get('first_name', '')} {user.get('last_name', '')}")
            if user.get('email'):
                context_parts.append(f"Email: {user.get('email')}")
        
        # Get family members if allowed
        if settings.allow_family_coordination:
            # Get user's household
            household = await self.db.households.find_one(
                {"members": {"$elemMatch": {"user_id": user_id}}},
                {"_id": 0}
            )
            if household:
                member_ids = [m.get('user_id') for m in household.get('members', [])]
                family_members = await self.db.users.find(
                    {"id": {"$in": member_ids}},
                    {"_id": 0, "first_name": 1, "last_name": 1, "id": 1}
                ).to_list(100)
                if family_members:
                    names = [f"{m.get('first_name', '')} {m.get('last_name', '')}" for m in family_members if m.get('id') != user_id]
                    if names:
                        context_parts.append(f"Члены семьи: {', '.join(names)}")
        
        # Get financial summary if allowed
        if settings.allow_financial_analysis:
            # Get recent transactions
            transactions = await self.db.transactions.find(
                {"user_id": user_id},
                {"_id": 0}
            ).sort("created_at", -1).limit(10).to_list(10)
            
            if transactions:
                total_income = sum(t.get('amount', 0) for t in transactions if t.get('type') == 'income')
                total_expense = sum(t.get('amount', 0) for t in transactions if t.get('type') == 'expense')
                context_parts.append(f"Финансы (последние транзакции): Доход: {total_income}, Расходы: {total_expense}")
        
        # Get upcoming events
        upcoming_events = await self.db.events.find(
            {"$or": [{"user_id": user_id}, {"participants": user_id}]},
            {"_id": 0, "title": 1, "date": 1}
        ).sort("date", 1).limit(5).to_list(5)
        
        if upcoming_events:
            events_str = "; ".join([f"{e.get('title', 'Событие')} ({e.get('date', '')})" for e in upcoming_events])
            context_parts.append(f"Предстоящие события: {events_str}")
        
        return "\n".join(context_parts) if context_parts else "Контекст недоступен"
    
    async def chat(self, user_id: str, request: ChatRequest) -> ChatResponse:
        """Process chat message and return AI response"""
        
        if not deepseek_client:
            raise Exception("DeepSeek API not configured")
        
        # Get user settings
        settings_doc = await self.db.agent_settings.find_one(
            {"user_id": user_id},
            {"_id": 0}
        )
        settings = AgentSettings(**settings_doc) if settings_doc else AgentSettings(user_id=user_id)
        
        # Get or create conversation
        conversation = await self.get_or_create_conversation(user_id, request.conversation_id)
        
        # Build user context
        user_context = await self.get_user_context(user_id, settings)
        
        # Add user message to conversation
        user_message = AgentMessage(role="user", content=request.message)
        conversation.messages.append(user_message)
        
        # Build messages for API call
        messages = [
            {"role": "system", "content": f"{ERIC_SYSTEM_PROMPT}\n\n## Текущий контекст пользователя:\n{user_context}"}
        ]
        
        # Add conversation history (last 10 messages for context)
        for msg in conversation.messages[-10:]:
            if msg.role in ['user', 'assistant']:
                messages.append({"role": msg.role, "content": msg.content})
        
        try:
            # Call DeepSeek API
            response = await deepseek_client.chat.completions.create(
                model=self.model,
                messages=messages,
                max_tokens=2000,
                temperature=0.7
            )
            
            assistant_content = response.choices[0].message.content
            
            # Create assistant message
            assistant_message = AgentMessage(
                role="assistant",
                content=assistant_content
            )
            conversation.messages.append(assistant_message)
            
            # Update conversation title if it's the first exchange
            if len(conversation.messages) == 2:
                # Use first ~50 chars of user message as title
                conversation.title = request.message[:50] + ("..." if len(request.message) > 50 else "")
            
            # Update conversation in database
            conversation.updated_at = datetime.now(timezone.utc).isoformat()
            await self.db.agent_conversations.update_one(
                {"id": conversation.id},
                {"$set": conversation.dict()},
                upsert=True
            )
            
            return ChatResponse(
                conversation_id=conversation.id,
                message=assistant_message,
                suggested_actions=[]
            )
            
        except Exception as e:
            # Return error message
            error_message = AgentMessage(
                role="assistant",
                content=f"Извините, произошла ошибка при обработке запроса. Пожалуйста, попробуйте позже. 😔\n\nОшибка: {str(e)}"
            )
            return ChatResponse(
                conversation_id=conversation.id,
                message=error_message,
                suggested_actions=[]
            )
    
    async def get_conversations(self, user_id: str, limit: int = 20, offset: int = 0) -> List[Dict]:
        """Get user's conversation history"""
        conversations = await self.db.agent_conversations.find(
            {"user_id": user_id},
            {"_id": 0, "messages": 0}  # Exclude messages for list view
        ).sort("updated_at", -1).skip(offset).limit(limit).to_list(limit)
        
        return conversations
    
    async def get_conversation(self, user_id: str, conversation_id: str) -> Optional[Dict]:
        """Get specific conversation with messages"""
        conversation = await self.db.agent_conversations.find_one(
            {"id": conversation_id, "user_id": user_id},
            {"_id": 0}
        )
        return conversation
    
    async def delete_conversation(self, user_id: str, conversation_id: str) -> bool:
        """Delete a conversation"""
        result = await self.db.agent_conversations.delete_one(
            {"id": conversation_id, "user_id": user_id}
        )
        return result.deleted_count > 0
    
    async def get_settings(self, user_id: str) -> AgentSettings:
        """Get user's agent settings"""
        settings_doc = await self.db.agent_settings.find_one(
            {"user_id": user_id},
            {"_id": 0}
        )
        if settings_doc:
            return AgentSettings(**settings_doc)
        
        # Create default settings
        default_settings = AgentSettings(user_id=user_id)
        await self.db.agent_settings.insert_one(default_settings.dict())
        return default_settings
    
    async def update_settings(self, user_id: str, updates: Dict[str, Any]) -> AgentSettings:
        """Update user's agent settings"""
        updates["updated_at"] = datetime.now(timezone.utc).isoformat()
        
        await self.db.agent_settings.update_one(
            {"user_id": user_id},
            {"$set": updates},
            upsert=True
        )
        
        return await self.get_settings(user_id)

    async def process_post_mention(self, user_id: str, post_id: str, post_content: str, author_name: str) -> str:
        """Process @ERIC mention in a post and generate a comment response"""
        
        if not deepseek_client:
            return "ERIC не настроен. Пожалуйста, обратитесь к администратору."
        
        # Get user settings
        settings_doc = await self.db.agent_settings.find_one(
            {"user_id": user_id},
            {"_id": 0}
        )
        settings = AgentSettings(**settings_doc) if settings_doc else AgentSettings(user_id=user_id)
        
        # Build context
        user_context = await self.get_user_context(user_id, settings)
        
        # Special prompt for post mentions
        post_prompt = f"""Пользователь {author_name} упомянул тебя в посте. Ответь как комментарий к этому посту.

## Пост:
{post_content}

## Твоя задача:
- Проанализируй пост и ответь полезным комментарием
- Если просят что-то найти/проанализировать - сделай это на основе данных платформы
- Будь краток и полезен (это комментарий, не длинный диалог)
- Используй дружелюбный тон

## Контекст пользователя:
{user_context}"""

        try:
            response = await deepseek_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": ERIC_SYSTEM_PROMPT},
                    {"role": "user", "content": post_prompt}
                ],
                max_tokens=500,
                temperature=0.7
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            return f"Извините, не смог обработать запрос: {str(e)}"

    async def analyze_image(self, user_id: str, image_base64: str, mime_type: str, question: Optional[str] = None) -> Dict[str, Any]:
        """
        Analyze an image using Claude Sonnet 4.5 via Emergent LLM Key
        
        Args:
            user_id: The user's ID
            image_base64: Base64 encoded image data
            mime_type: MIME type of the image (image/jpeg, image/png, image/webp)
            question: Optional specific question about the image
        
        Returns:
            Dict with analysis results
        """
        if not EMERGENT_LLM_KEY:
            return {
                "success": False,
                "error": "Анализ изображений недоступен. Ключ API не настроен."
            }
        
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent
            
            # Create a new chat instance for this analysis
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=f"image-analysis-{user_id}-{uuid4()}",
                system_message="""Ты ERIC - ИИ-помощник платформы ZION.CITY. 
Ты анализируешь изображения и предоставляешь полезную информацию на русском языке.
Будь дружелюбным и полезным. Отвечай структурированно."""
            )
            
            # Configure to use Claude Sonnet for vision
            chat.with_model("anthropic", "claude-sonnet-4-5-20250929")
            
            # Create image content
            image_content = ImageContent(image_base64=image_base64)
            
            # Build the question
            analysis_question = question if question else "Опиши подробно, что изображено на этой картинке. Укажи ключевые объекты, детали и контекст."
            
            # Create message with image
            user_message = UserMessage(
                text=analysis_question,
                file_contents=[image_content]
            )
            
            # Send and get response
            response = await chat.send_message(user_message)
            
            return {
                "success": True,
                "analysis": response,
                "question": analysis_question
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Ошибка при анализе изображения: {str(e)}"
            }

    async def analyze_document(self, user_id: str, document_text: str, document_name: str, question: Optional[str] = None) -> Dict[str, Any]:
        """
        Analyze a document using DeepSeek (cost-optimized).
        DeepSeek handles all text-based documents: PDF, DOCX, TXT, CSV, XLSX, JSON, XML, etc.
        Claude is reserved for image analysis only.
        
        Args:
            user_id: The user's ID
            document_text: The extracted text content of the document
            document_name: Name of the document
            question: Optional specific question about the document
        
        Returns:
            Dict with analysis results
        """
        if not deepseek_client:
            return {
                "success": False,
                "error": "Анализ документов недоступен. DeepSeek API не настроен."
            }
        
        try:
            # System prompt for document analysis
            system_prompt = """Ты ERIC - ИИ-помощник платформы ZION.CITY. 
Ты анализируешь документы и предоставляешь полезную информацию на русском языке.
Будь дружелюбным и полезным. Отвечай структурированно.
При анализе документов:
- Выдели ключевые моменты
- Если это финансовый документ - обрати внимание на суммы и даты
- Если это договор - укажи важные условия
- Если это таблица/CSV/Excel - проанализируй данные и тренды
- Предложи действия если уместно"""
            
            # Build the prompt
            if question:
                prompt = f"""Документ: {document_name}

Содержимое документа:
---
{document_text[:30000]}
---

Вопрос пользователя: {question}"""
            else:
                prompt = f"""Документ: {document_name}

Содержимое документа:
---
{document_text[:30000]}
---

Проанализируй этот документ и предоставь:
1. Краткое содержание (2-3 предложения)
2. Ключевые моменты
3. Важные даты и суммы (если есть)
4. Рекомендации или действия (если уместно)"""
            
            # Call DeepSeek API
            response = await deepseek_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.5  # Lower temperature for more factual document analysis
            )
            
            analysis_content = response.choices[0].message.content
            
            return {
                "success": True,
                "analysis": analysis_content,
                "document_name": document_name,
                "question": question,
                "model_used": "deepseek"  # Track which model was used
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Ошибка при анализе документа: {str(e)}"
            }
    
    async def analyze_file_smart(self, user_id: str, file_content: bytes, filename: str, mime_type: str = None, question: Optional[str] = None) -> Dict[str, Any]:
        """
        Smart file analysis that routes to the appropriate LLM based on file type.
        - Images -> Claude Sonnet (vision required)
        - Documents/Text -> DeepSeek (cheaper)
        
        Args:
            user_id: The user's ID
            file_content: Raw bytes of the file
            filename: Original filename
            mime_type: MIME type of the file
            question: Optional question about the file
        
        Returns:
            Dict with analysis results and routing info
        """
        # Detect file type and routing
        file_category, use_deepseek = detect_file_type(filename, mime_type)
        
        if not use_deepseek:
            # Route to Claude for images
            image_base64 = base64.b64encode(file_content).decode('utf-8')
            result = await self.analyze_image(
                user_id=user_id,
                image_base64=image_base64,
                mime_type=mime_type or 'image/png',
                question=question
            )
            result["routing"] = {
                "file_category": file_category,
                "model_used": "claude-sonnet",
                "reason": "Image analysis requires vision capabilities"
            }
            return result
        
        # Route to DeepSeek for documents
        extracted_text, file_type_desc = extract_text_from_file(file_content, filename, mime_type)
        
        result = await self.analyze_document(
            user_id=user_id,
            document_text=extracted_text,
            document_name=f"{filename} ({file_type_desc})",
            question=question
        )
        result["routing"] = {
            "file_category": file_category,
            "model_used": "deepseek",
            "reason": "Text-based document analysis - cost optimized"
        }
        return result

    async def chat_with_image(self, user_id: str, message: str, image_base64: str, mime_type: str, conversation_id: Optional[str] = None) -> ChatResponse:
        """
        Chat with ERIC while providing an image for context
        Uses Claude Sonnet 4.5 for vision capabilities
        """
        if not EMERGENT_LLM_KEY:
            error_message = AgentMessage(
                role="assistant",
                content="Извините, анализ изображений временно недоступен. Пожалуйста, попробуйте позже."
            )
            return ChatResponse(
                conversation_id=conversation_id or str(uuid4()),
                message=error_message,
                suggested_actions=[]
            )
        
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent
            
            # Get user settings
            settings_doc = await self.db.agent_settings.find_one(
                {"user_id": user_id},
                {"_id": 0}
            )
            settings = AgentSettings(**settings_doc) if settings_doc else AgentSettings(user_id=user_id)
            
            # Get or create conversation
            conversation = await self.get_or_create_conversation(user_id, conversation_id)
            
            # Build user context
            user_context = await self.get_user_context(user_id, settings)
            
            # Add user message to conversation (without image for storage)
            user_message_record = AgentMessage(role="user", content=f"[Изображение] {message}")
            conversation.messages.append(user_message_record)
            
            # Create chat instance with Claude
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=f"chat-image-{user_id}-{conversation.id}",
                system_message=f"""{ERIC_SYSTEM_PROMPT}

## Текущий контекст пользователя:
{user_context}

## Важно:
Пользователь прикрепил изображение к сообщению. Проанализируй его и ответь на вопрос пользователя."""
            )
            
            chat.with_model("anthropic", "claude-sonnet-4-5-20250929")
            
            # Create image content and message
            image_content = ImageContent(image_base64=image_base64)
            user_msg = UserMessage(
                text=message,
                file_contents=[image_content]
            )
            
            # Get response
            response_text = await chat.send_message(user_msg)
            
            # Create assistant message
            assistant_message = AgentMessage(
                role="assistant",
                content=response_text
            )
            conversation.messages.append(assistant_message)
            
            # Update conversation title if first exchange
            if len(conversation.messages) == 2:
                conversation.title = f"📷 {message[:40]}..." if len(message) > 40 else f"📷 {message}"
            
            # Save conversation
            conversation.updated_at = datetime.now(timezone.utc).isoformat()
            await self.db.agent_conversations.update_one(
                {"id": conversation.id},
                {"$set": conversation.dict()},
                upsert=True
            )
            
            return ChatResponse(
                conversation_id=conversation.id,
                message=assistant_message,
                suggested_actions=[]
            )
            
        except Exception as e:
            error_message = AgentMessage(
                role="assistant",
                content=f"Извините, произошла ошибка при анализе изображения: {str(e)}"
            )
            return ChatResponse(
                conversation_id=conversation_id or str(uuid4()),
                message=error_message,
                suggested_actions=[]
            )

    # ===== SEARCH FUNCTIONALITY =====
    
    async def search_platform(self, user_id: str, query: str, search_type: str = "all", location: Optional[str] = None, limit: int = 10) -> Dict[str, Any]:
        """
        Search across the ZION.CITY platform for services, products, people, and organizations.
        Uses both database search and AI-powered relevance ranking.
        """
        results = {
            "query": query,
            "search_type": search_type,
            "results": [],
            "ai_summary": ""
        }
        
        try:
            search_results = []
            query_lower = query.lower()
            
            # Search Organizations/Businesses
            if search_type in ["all", "organizations"]:
                orgs = await self.db.work_organizations.find({
                    "$or": [
                        {"name": {"$regex": query, "$options": "i"}},
                        {"description": {"$regex": query, "$options": "i"}},
                        {"industry": {"$regex": query, "$options": "i"}},
                        {"address_city": {"$regex": query, "$options": "i"}}
                    ]
                }).limit(limit).to_list(limit)
                
                for org in orgs:
                    # Skip private organizations
                    if org.get("is_private", False):
                        continue
                    search_results.append({
                        "id": org.get("id") or org.get("organization_id"),
                        "type": "organization",
                        "name": org.get("name"),
                        "description": org.get("description"),
                        "industry": org.get("industry"),
                        "metadata": {
                            "member_count": org.get("member_count", 0),
                            "founded_year": org.get("founded_year"),
                            "logo_url": org.get("logo_url"),
                            "city": org.get("address_city"),
                            "organization_type": org.get("organization_type")
                        }
                    })
            
            # Search Services (service_listings collection)
            if search_type in ["all", "services"]:
                services = await self.db.service_listings.find({
                    "$or": [
                        {"name": {"$regex": query, "$options": "i"}},
                        {"description": {"$regex": query, "$options": "i"}},
                        {"category_id": {"$regex": query, "$options": "i"}},
                        {"subcategory_id": {"$regex": query, "$options": "i"}},
                        {"city": {"$regex": query, "$options": "i"}}
                    ],
                    "status": "ACTIVE"
                }).limit(limit).to_list(limit)
                
                for svc in services:
                    search_results.append({
                        "id": svc.get("id"),
                        "type": "service",
                        "name": svc.get("name"),
                        "description": svc.get("description"),
                        "metadata": {
                            "category": svc.get("category_id"),
                            "subcategory": svc.get("subcategory_id"),
                            "price_from": svc.get("price_from"),
                            "price_to": svc.get("price_to"),
                            "price_type": svc.get("price_type"),
                            "rating": svc.get("rating"),
                            "review_count": svc.get("review_count", 0),
                            "city": svc.get("city"),
                            "currency": svc.get("currency", "RUB")
                        }
                    })
            
            # Search Products in Marketplace (marketplace_products collection)
            if search_type in ["all", "products"]:
                products = await self.db.marketplace_products.find({
                    "$or": [
                        {"title": {"$regex": query, "$options": "i"}},
                        {"description": {"$regex": query, "$options": "i"}},
                        {"category": {"$regex": query, "$options": "i"}}
                    ],
                    "status": {"$in": ["available", "AVAILABLE"]}
                }).limit(limit).to_list(limit)
                
                for prod in products:
                    search_results.append({
                        "id": prod.get("id"),
                        "type": "product",
                        "name": prod.get("title"),
                        "description": prod.get("description"),
                        "metadata": {
                            "price": prod.get("price"),
                            "currency": prod.get("currency", "RUB"),
                            "condition": prod.get("condition"),
                            "category": prod.get("category"),
                            "city": prod.get("city"),
                            "accept_altyn": prod.get("accept_altyn", False),
                            "altyn_price": prod.get("altyn_price")
                        }
                    })
            
            # Search People (public profiles only)
            if search_type in ["all", "people"]:
                users = await self.db.users.find({
                    "$or": [
                        {"first_name": {"$regex": query, "$options": "i"}},
                        {"last_name": {"$regex": query, "$options": "i"}},
                        {"bio": {"$regex": query, "$options": "i"}}
                    ]
                }, {"_id": 0, "password_hash": 0}).limit(limit).to_list(limit)
                
                for user in users:
                    # Check privacy settings
                    privacy = user.get("privacy_settings", {})
                    if privacy.get("profile_visibility", "public") == "public":
                        search_results.append({
                            "id": user.get("id"),
                            "type": "person",
                            "name": f"{user.get('first_name', '')} {user.get('last_name', '')}".strip(),
                            "description": user.get("bio"),
                            "metadata": {
                                "profile_picture": user.get("profile_picture")
                            }
                        })
            
            # Apply limit to total results
            search_results = search_results[:limit]
            
            results["results"] = search_results
            results["total_count"] = len(search_results)
            
            # Generate AI summary if results found
            if search_results and DEEPSEEK_API_KEY:
                summary_prompt = f"""На основе результатов поиска по запросу "{query}", дай краткую рекомендацию пользователю.
                
Найдено {len(search_results)} результатов:
{[f"- {r['type']}: {r['name']}" for r in search_results[:5]]}

Ответь кратко (2-3 предложения) на русском языке."""

                try:
                    response = await deepseek_client.chat.completions.create(
                        model=self.model,
                        messages=[
                            {"role": "system", "content": "Ты помощник поиска ZION.CITY. Давай краткие и полезные рекомендации."},
                            {"role": "user", "content": summary_prompt}
                        ],
                        max_tokens=150,
                        temperature=0.7
                    )
                    results["ai_summary"] = response.choices[0].message.content
                except:
                    results["ai_summary"] = f"Найдено {len(search_results)} результатов по вашему запросу."
            
            return results
            
        except Exception as e:
            return {
                "query": query,
                "search_type": search_type,
                "results": [],
                "error": str(e)
            }

    # ===== BUSINESS ERIC SETTINGS =====
    
    async def get_business_settings(self, organization_id: str) -> Optional[BusinessERICSettings]:
        """Get ERIC settings for a business/organization"""
        settings_doc = await self.db.business_eric_settings.find_one(
            {"organization_id": organization_id},
            {"_id": 0}
        )
        if settings_doc:
            return BusinessERICSettings(**settings_doc)
        return None
    
    async def save_business_settings(self, settings: BusinessERICSettings) -> BusinessERICSettings:
        """Save or update ERIC settings for a business/organization"""
        settings.updated_at = datetime.now(timezone.utc).isoformat()
        await self.db.business_eric_settings.update_one(
            {"organization_id": settings.organization_id},
            {"$set": settings.dict()},
            upsert=True
        )
        return settings
    
    async def query_business_eric(self, user_id: str, organization_id: str, query: str) -> Dict[str, Any]:
        """
        Query a business's ERIC agent for information.
        Respects the business's privacy settings.
        """
        # Get business ERIC settings
        settings = await self.get_business_settings(organization_id)
        if not settings:
            # Create default settings
            settings = BusinessERICSettings(organization_id=organization_id)
            await self.save_business_settings(settings)
        
        # Check if queries are allowed
        if not settings.allow_user_eric_queries:
            return {
                "success": False,
                "error": "Этот бизнес не принимает запросы от ERIC-помощников"
            }
        
        # Get organization info
        org = await self.db.work_organizations.find_one(
            {"id": organization_id},
            {"_id": 0}
        )
        if not org:
            return {
                "success": False,
                "error": "Организация не найдена"
            }
        
        # Build response based on allowed data
        response_data = {}
        
        if settings.share_public_data:
            response_data["company_info"] = {
                "name": org.get("name"),
                "description": org.get("description"),
                "industry": org.get("industry"),
                "website": org.get("website"),
                "email": org.get("official_email")
            }
        
        if settings.share_promotions:
            # Get active promotions
            promos = await self.db.promotions.find({
                "organization_id": organization_id,
                "is_active": True
            }, {"_id": 0}).to_list(10)
            response_data["promotions"] = promos
        
        if settings.share_aggregated_analytics and settings.share_repeat_customer_stats:
            # This would be calculated from transaction data
            # For now, return placeholder
            response_data["analytics"] = {
                "repeat_customer_rate": "Данные недоступны",
                "note": "Агрегированная аналитика"
            }
        
        if settings.share_ratings_reviews:
            # Get average rating
            reviews = await self.db.reviews.find({
                "organization_id": organization_id
            }).to_list(100)
            if reviews:
                avg_rating = sum(r.get("rating", 0) for r in reviews) / len(reviews)
                response_data["ratings"] = {
                    "average_rating": round(avg_rating, 1),
                    "review_count": len(reviews)
                }
        
        return {
            "success": True,
            "organization_id": organization_id,
            "organization_name": org.get("name"),
            "data": response_data,
            "settings_active": settings.is_active
        }

    async def query_multiple_businesses(self, user_id: str, query: str, category: str = None, limit: int = 5) -> Dict[str, Any]:
        """
        Query multiple business ERICs for recommendations.
        Used when user asks for best/recommended services.
        Returns aggregated responses respecting each business's privacy settings.
        """
        results = []
        
        # Find businesses that allow queries
        query_filter = {}
        if category:
            # Map Russian category keywords to fields
            category_fields = {
                "красота": ["beauty", "салон"],
                "ремонт": ["repair", "сервис"],
                "машина": ["auto", "автосервис"],
                "еда": ["food", "ресторан", "кафе"],
                "здоровье": ["health", "медицина"],
                "образование": ["education", "школа"]
            }
            search_terms = category_fields.get(category.lower(), [category])
            query_filter["$or"] = [
                {"industry": {"$regex": term, "$options": "i"}} for term in search_terms
            ] + [
                {"name": {"$regex": term, "$options": "i"}} for term in search_terms
            ]
        
        # Get organizations
        orgs = await self.db.work_organizations.find(
            query_filter, {"_id": 0}
        ).limit(limit * 2).to_list(limit * 2)  # Get more to filter by settings
        
        for org in orgs:
            org_id = org.get("id") or org.get("organization_id")
            if not org_id:
                continue
                
            # Get business ERIC settings
            settings = await self.get_business_settings(org_id)
            
            # Default: allow queries if no settings exist
            allow_queries = True
            if settings:
                allow_queries = settings.allow_user_eric_queries
            
            if not allow_queries:
                continue
            
            # Query this business's ERIC
            business_response = await self.query_business_eric(user_id, org_id, query)
            
            if business_response.get("success"):
                # Calculate a simple relevance score
                score = 0
                data = business_response.get("data", {})
                
                # Boost score based on available data
                if data.get("company_info"):
                    score += 1
                if data.get("ratings"):
                    rating = data["ratings"].get("average_rating", 0)
                    score += rating / 5  # Normalize to 0-1
                if data.get("promotions"):
                    score += 0.5
                
                results.append({
                    "organization_id": org_id,
                    "organization_name": business_response.get("organization_name"),
                    "data": data,
                    "relevance_score": score,
                    "industry": org.get("industry"),
                    "city": org.get("address_city")
                })
            
            if len(results) >= limit:
                break
        
        # Sort by relevance score
        results.sort(key=lambda x: x.get("relevance_score", 0), reverse=True)
        
        # Create notification if we got results
        final_results = results[:limit]
        if final_results:
            # Create a summary of recommendations for the notification
            business_names = [r.get("organization_name", "Бизнес") for r in final_results[:3]]
            notification_message = f"Найдено {len(final_results)} рекомендаций по запросу «{query}»: {', '.join(business_names)}"
            if len(final_results) > 3:
                notification_message += f" и еще {len(final_results) - 3}"
            
            await self.create_notification(
                user_id=user_id,
                notification_type="eric_recommendation",
                title="🎯 ERIC нашёл рекомендации",
                message=notification_message,
                related_data={
                    "query": query,
                    "category": category,
                    "results_count": len(final_results),
                    "business_ids": [r.get("organization_id") for r in final_results]
                }
            )
        
        return {
            "query": query,
            "category": category,
            "results": final_results,
            "total_businesses_queried": len(orgs),
            "businesses_responding": len(results)
        }

    async def chat_with_search(self, user_id: str, message: str, conversation_id: Optional[str] = None) -> ChatResponse:
        """
        Enhanced chat that can perform platform searches when needed.
        ERIC will automatically search when user asks about finding services, products, people.
        """
        # Keywords that trigger search
        search_keywords = ["найди", "найти", "поиск", "ищу", "где", "какой", "какая", "лучший", "лучшая", "рекомендуй", "посоветуй", "покажи"]
        should_search = any(kw in message.lower() for kw in search_keywords)
        
        # Extract actual search query by removing trigger words
        search_query = message.lower()
        for kw in search_keywords:
            search_query = search_query.replace(kw, "")
        search_query = search_query.strip()
        
        # Map Russian keywords to English category/search terms
        # Include word stems and common forms
        category_mappings = {
            "красот": ["beauty", "салон", "маникюр", "педикюр", "стрижка", "парикмахер"],
            "ремонт": ["repair", "сервис", "мастер", "техника"],
            "машин": ["auto", "car", "автосервис", "шиномонтаж"],
            "автосервис": ["auto", "car", "ремонт", "шиномонтаж"],
            "еда": ["food", "ресторан", "кафе", "доставка"],
            "еду": ["food", "ресторан", "кафе", "доставка"],
            "здоров": ["health", "медицина", "врач", "клиника"],
            "образован": ["education", "школа", "курсы", "репетитор"],
            "школ": ["education", "школа", "курсы", "обучение"],
            "услуг": ["service", "услуга", "сервис", "тест"],
            "товар": ["product", "товар", "магазин"],
            "люд": ["person", "человек"],
            "организац": ["organization", "компания", "фирма"],
            "тест": ["тест", "test", "услуга"]
        }
        
        # Check if user is asking for recommendations (triggers inter-agent queries)
        recommendation_keywords = ["лучший", "лучшая", "лучшее", "рекомендуй", "посоветуй", "какой лучше", "где лучше"]
        wants_recommendations = any(kw in message.lower() for kw in recommendation_keywords)
        
        # Detect category for business queries
        detected_category = None
        for ru_term in category_mappings.keys():
            if ru_term in search_query:
                detected_category = ru_term
                break
        
        # Expand search query with mapped terms
        expanded_terms = [search_query]
        for ru_term, en_terms in category_mappings.items():
            if ru_term in search_query:
                expanded_terms.extend(en_terms)
        
        search_context = ""
        found_results = False
        action_cards = []  # Initialize action cards
        business_recommendations = []  # For inter-agent results
        
        if should_search:
            # Try searching with each expanded term until we find results
            all_results = []
            for term in expanded_terms:
                if term:
                    search_result = await self.search_platform(user_id, term, "all", limit=5)
                    if search_result.get("results"):
                        for r in search_result['results']:
                            # Avoid duplicates
                            if not any(existing['id'] == r['id'] for existing in all_results):
                                all_results.append(r)
                    if len(all_results) >= 5:
                        break
            
            # If user wants recommendations, also query business ERICs
            if wants_recommendations and detected_category:
                business_query_result = await self.query_multiple_businesses(
                    user_id=user_id,
                    query=message,
                    category=detected_category,
                    limit=3
                )
                business_recommendations = business_query_result.get("results", [])
            
            if all_results:
                found_results = True
                results_formatted = []
                
                for r in all_results[:5]:
                    result_str = f"- **{r['type'].upper()}**: {r['name']}"
                    if r.get('description'):
                        result_str += f" - {r['description'][:100]}"
                    if r.get('metadata'):
                        meta = r['metadata']
                        if meta.get('price_from'):
                            result_str += f" | Цена от: {meta['price_from']} {meta.get('currency', 'RUB')}"
                        if meta.get('city'):
                            result_str += f" | Город: {meta['city']}"
                        if meta.get('rating'):
                            result_str += f" | Рейтинг: {meta['rating']}⭐"
                    results_formatted.append(result_str)
                    
                    # Create action card for each result
                    action_card = {
                        "id": r.get('id'),
                        "type": r.get('type'),
                        "name": r.get('name'),
                        "description": r.get('description', '')[:100] if r.get('description') else '',
                        "metadata": r.get('metadata', {})
                    }
                    
                    # Add navigation info based on type
                    if r['type'] == 'service':
                        action_card["action"] = {
                            "label": "Забронировать",
                            "icon": "calendar",
                            "route": f"/services/{r['id']}",
                            "type": "navigate"
                        }
                    elif r['type'] == 'organization':
                        action_card["action"] = {
                            "label": "Подробнее",
                            "icon": "building",
                            "route": f"/organizations/{r['id']}",
                            "type": "navigate"
                        }
                    elif r['type'] == 'product':
                        action_card["action"] = {
                            "label": "Посмотреть",
                            "icon": "shopping-bag",
                            "route": f"/marketplace/{r['id']}",
                            "type": "navigate"
                        }
                    elif r['type'] == 'person':
                        action_card["action"] = {
                            "label": "Написать",
                            "icon": "message",
                            "route": f"/messages?user={r['id']}",
                            "type": "navigate"
                        }
                    
                    action_cards.append(action_card)
                
                search_context = f"""
## РЕЗУЛЬТАТЫ ПОИСКА ПО ПЛАТФОРМЕ ZION.CITY (НАЙДЕНО {len(all_results)} результатов):
{chr(10).join(results_formatted)}

ВАЖНО: Ты ДОЛЖЕН использовать эти результаты в своём ответе. Представь их кратко и дружелюбно. Пользователь увидит интерактивные карточки с кнопками действий.
"""
            else:
                # NO RESULTS FOUND - Provide honest response without hallucination
                search_context = f"""
## ПОИСК ПО ПЛАТФОРМЕ ZION.CITY: НИЧЕГО НЕ НАЙДЕНО

Поисковый запрос пользователя: "{message}"
Результат поиска в базе данных: 0 результатов

КРИТИЧЕСКИ ВАЖНО - НЕ ВЫДУМЫВАЙ РЕЗУЛЬТАТЫ!
Ты ОБЯЗАН честно сказать пользователю, что по его запросу ничего не найдено в базе данных ZION.CITY.

Используй ТОЧНО этот шаблон ответа (можешь немного адаптировать под контекст запроса):

---
К сожалению, по вашему запросу пока ничего не найдено на платформе ZION.CITY. 😔

ZION.CITY — это молодая и растущая платформа, и возможно в вашем регионе ещё не зарегистрированы поставщики таких услуг или товаров.

💡 **Что можно сделать:**
- Попробуйте изменить или расширить поисковый запрос
- Проверьте раздел «Сервисы» или «Организации» — возможно, там есть похожие предложения
- Расскажите друзьям и знакомым о ZION.CITY — платформа **полностью бесплатна** для малого бизнеса и частных специалистов!

Чем больше людей узнают о ZION.CITY, тем больше услуг и товаров появится в вашем регионе! 🚀
---

НЕ ДОБАВЛЯЙ никаких выдуманных организаций, услуг или контактов!
"""
            
            # Add business recommendations from inter-agent queries
            if business_recommendations:
                business_context = "\n## РЕКОМЕНДАЦИИ ОТ БИЗНЕС-ERIC (запросы к бизнес-помощникам):\n"
                for br in business_recommendations:
                    org_name = br.get("organization_name", "Неизвестно")
                    data = br.get("data", {})
                    business_context += f"\n### {org_name}"
                    if br.get("city"):
                        business_context += f" ({br['city']})"
                    business_context += "\n"
                    
                    if data.get("company_info"):
                        info = data["company_info"]
                        if info.get("description"):
                            business_context += f"- {info['description'][:150]}\n"
                    
                    if data.get("ratings"):
                        ratings = data["ratings"]
                        business_context += f"- Рейтинг: {ratings.get('average_rating', 'N/A')}⭐ ({ratings.get('review_count', 0)} отзывов)\n"
                    
                    if data.get("promotions") and len(data["promotions"]) > 0:
                        business_context += f"- 🎁 Есть активные акции!\n"
                    
                    # Add action card for this business
                    action_cards.append({
                        "id": br.get("organization_id"),
                        "type": "recommendation",
                        "name": org_name,
                        "description": data.get("company_info", {}).get("description", "")[:100] if data.get("company_info") else "",
                        "metadata": {
                            "city": br.get("city"),
                            "rating": data.get("ratings", {}).get("average_rating"),
                            "industry": br.get("industry"),
                            "has_promotions": bool(data.get("promotions"))
                        },
                        "action": {
                            "label": "Подробнее",
                            "icon": "star",
                            "route": f"/organizations/{br.get('organization_id')}",
                            "type": "navigate"
                        }
                    })
                
                search_context += business_context
        
        # Get user settings
        settings_doc = await self.db.agent_settings.find_one(
            {"user_id": user_id},
            {"_id": 0}
        )
        settings = AgentSettings(**settings_doc) if settings_doc else AgentSettings(user_id=user_id)
        
        # Get or create conversation
        conversation = await self.get_or_create_conversation(user_id, conversation_id)
        
        # Add user message
        user_message = AgentMessage(role="user", content=message)
        conversation.messages.append(user_message)
        
        # Build context
        user_context = await self.get_user_context(user_id, settings)
        
        # Enhanced system prompt with search context
        enhanced_prompt = ERIC_SYSTEM_PROMPT + f"""

## Дополнительный контекст пользователя:
{user_context}

{search_context}

## КРИТИЧЕСКИЕ ИНСТРУКЦИИ ПО ПОИСКУ:
- Ты работаешь ТОЛЬКО с данными из базы ZION.CITY
- НИКОГДА не выдумывай организации, услуги, товары или контакты
- Если выше есть "РЕЗУЛЬТАТЫ ПОИСКА" - используй ТОЛЬКО эти результаты
- Если написано "НИЧЕГО НЕ НАЙДЕНО" - честно скажи об этом пользователю
- НЕ предлагай искать в интернете или на других платформах
- Продвигай ZION.CITY как бесплатную платформу для малого бизнеса
"""
        
        # Build messages for API
        api_messages = [{"role": "system", "content": enhanced_prompt}]
        for msg in conversation.messages[-10:]:
            api_messages.append({"role": msg.role, "content": msg.content})
        
        try:
            response = await deepseek_client.chat.completions.create(
                model=self.model,
                messages=api_messages,
                max_tokens=1000,
                temperature=0.7
            )
            
            assistant_content = response.choices[0].message.content
            assistant_message = AgentMessage(role="assistant", content=assistant_content)
            conversation.messages.append(assistant_message)
            
            # Update conversation
            if len(conversation.messages) == 2:
                conversation.title = message[:50] + "..." if len(message) > 50 else message
            
            conversation.updated_at = datetime.now(timezone.utc).isoformat()
            await self.db.agent_conversations.update_one(
                {"id": conversation.id},
                {"$set": conversation.dict()},
                upsert=True
            )
            
            # Build suggested actions from search results
            suggested_actions = []
            if action_cards:
                suggested_actions = [{
                    "type": "search_results",
                    "cards": action_cards
                }]
            
            return ChatResponse(
                conversation_id=conversation.id,
                message=assistant_message,
                suggested_actions=suggested_actions
            )
            
        except Exception as e:
            error_message = AgentMessage(
                role="assistant",
                content=f"Извините, произошла ошибка: {str(e)}"
            )
            return ChatResponse(
                conversation_id=conversation_id or str(uuid4()),
                message=error_message,
                suggested_actions=[]
            )
